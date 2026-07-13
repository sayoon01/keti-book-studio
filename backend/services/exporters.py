"""책 전체를 Markdown/DOCX/PDF로 렌더링한다.

PDF는 xhtml2pdf(순수 파이썬, reportlab 기반)를 쓴다 — weasyprint 같은
시스템 라이브러리(cairo/pango) 의존성이 없어 서버 환경을 덜 탄다.
단, 한글 렌더링을 위해서는 트루타입(glyf) 방식 한글 폰트가 필요하다
(나눔고딕 등 — Noto CJK는 OpenType CFF 방식이라 reportlab이 못 읽는다).
"""

import os
from pathlib import Path

CANDIDATE_KOREAN_FONT_PATHS = [
    os.environ.get("KETI_PDF_FONT_PATH", ""),
    "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
    "/usr/share/fonts/truetype/nanum/NanumMyeongjo.ttf",
]


def _find_korean_font() -> str:
    for path in CANDIDATE_KOREAN_FONT_PATHS:
        if path and Path(path).exists():
            return path
    raise RuntimeError(
        "PDF 생성용 한글 폰트를 찾을 수 없습니다. "
        "다음 명령으로 나눔폰트를 설치해주세요: sudo apt-get install -y fonts-nanum "
        "(다른 경로의 폰트를 쓰려면 KETI_PDF_FONT_PATH 환경변수로 지정하세요)"
    )


def render_markdown(book_title: str, units: list[dict], *, include_toc: bool = True) -> str:
    parts = [f"# {book_title}\n"]

    if include_toc:
        parts.append("## 목차\n")
        for u in units:
            parts.append(f"- {u['title']}")
        parts.append("")

    for u in units:
        parts.append(f"# {u['title']}\n")
        parts.append(u.get("body_md") or "*(아직 작성되지 않음)*")
        parts.append("")

    return "\n".join(parts)


def _markdown_to_html_body(markdown_text: str) -> str:
    lines = markdown_text.split("\n")
    html_parts: list[str] = []
    in_list = False

    def _close_list():
        nonlocal in_list
        if in_list:
            html_parts.append("</ul>")
            in_list = False

    for line in lines:
        stripped = line.strip()
        if not stripped:
            _close_list()
            continue
        if stripped.startswith("### "):
            _close_list()
            html_parts.append(f"<h3>{stripped[4:]}</h3>")
        elif stripped.startswith("## "):
            _close_list()
            html_parts.append(f"<h2>{stripped[3:]}</h2>")
        elif stripped.startswith("# "):
            _close_list()
            html_parts.append(f"<h1>{stripped[2:]}</h1>")
        elif stripped.startswith("- "):
            if not in_list:
                html_parts.append("<ul>")
                in_list = True
            html_parts.append(f"<li>{stripped[2:]}</li>")
        else:
            _close_list()
            html_parts.append(f"<p>{stripped}</p>")
    _close_list()
    return "\n".join(html_parts)


def render_docx(markdown_text: str, output_path: str, title: str) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=0)

    for line in markdown_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            doc.add_paragraph(stripped)

    doc.save(output_path)


def render_pdf(markdown_text: str, output_path: str, title: str) -> None:
    from xhtml2pdf import pisa

    font_path = _find_korean_font()
    body_html = _markdown_to_html_body(markdown_text)

    html = f"""<html>
<head>
<meta charset="utf-8">
<style>
@font-face {{
    font-family: KoreanFont;
    src: url({font_path});
}}
body {{ font-family: KoreanFont; font-size: 11pt; }}
h1 {{ font-size: 18pt; margin-top: 20pt; }}
h2 {{ font-size: 14pt; margin-top: 14pt; }}
h3 {{ font-size: 12pt; margin-top: 10pt; }}
</style>
</head>
<body>
{body_html}
</body>
</html>"""

    with open(output_path, "wb") as f:
        result = pisa.CreatePDF(html, dest=f)
    if result.err:
        raise RuntimeError(f"PDF 생성 중 오류가 발생했습니다 (err={result.err})")
