"""Phase 2 완료 조건 테스트.

1. txt/csv 업로드 -> SourceProfile 자동 생성 (LLM은 가짜 함수로 대체)
2. 동일 파일(내용 동일) 재업로드 시 새 source_id가 생기지 않고 기존 것을 반환
3. 목적(purpose)을 바꿔 재분석하면 SourceProfile이 갱신되고 row는 1개만 유지
4. extract_url 은 실제 네트워크 없이 주입된 http_get 함수로 파싱 로직만 검증
"""

import json
import os
import tempfile

import pytest
from fastapi.testclient import TestClient
from sqlmodel import SQLModel

os.environ["KETI_DB_PATH"] = os.path.join(tempfile.mkdtemp(), "test.db")
os.environ["KETI_UPLOAD_DIR"] = tempfile.mkdtemp()

from backend.services.llm_client import get_llm_call  # noqa: E402
from backend.main import app  # noqa: E402
from backend.services.extractors import extract_url  # noqa: E402
from backend.storage.database import engine  # noqa: E402


def _fake_llm_call(system_prompt: str, user_prompt: str) -> str:
    purpose_line = next(
        (line for line in user_prompt.splitlines() if line.startswith("분석 목적:")), ""
    )
    return json.dumps(
        {
            "summary": f"가짜 요약 ({purpose_line})",
            "main_topics": ["주제1", "주제2"],
            "key_findings": ["발견1"],
            "tables": [{"description": "표1", "usable_for_chart": True}],
            "limitations": ["실제 검증 데이터 없음"],
            "recommended_uses": ["기술서"],
        },
        ensure_ascii=False,
    )


@pytest.fixture(autouse=True)
def _reset_db():
    from backend.storage import models  # noqa: F401

    SQLModel.metadata.create_all(engine)
    app.dependency_overrides[get_llm_call] = lambda: _fake_llm_call
    yield
    app.dependency_overrides.clear()
    SQLModel.metadata.drop_all(engine)


@pytest.fixture
def client():
    return TestClient(app)


def _create_book(client) -> dict:
    resp = client.post("/api/books", json={"workspace_id": "ws-1", "title": "테스트 책"})
    assert resp.status_code == 200, resp.text
    return resp.json()


def test_upload_txt_and_analyze_creates_profile(client):
    book = _create_book(client)

    resp = client.post(
        f"/api/books/{book['book_id']}/sources/upload",
        files={"file": ("note.txt", b"ALD process temperature and pressure matter.", "text/plain")},
    )
    assert resp.status_code == 200, resp.text
    source = resp.json()
    assert source["source_type"] == "txt"
    assert source["status"] == "uploaded"

    analyze_resp = client.post(f"/api/sources/{source['source_id']}/analyze", json={})
    assert analyze_resp.status_code == 200, analyze_resp.text
    profile = analyze_resp.json()
    assert profile["main_topics"] == ["주제1", "주제2"]
    assert profile["source_id"] == source["source_id"]

    source_after = client.get(f"/api/books/{book['book_id']}/sources").json()[0]
    assert source_after["status"] == "analyzed"


def test_duplicate_upload_returns_existing_source(client):
    book = _create_book(client)
    content = b"same content twice"

    first = client.post(
        f"/api/books/{book['book_id']}/sources/upload",
        files={"file": ("a.txt", content, "text/plain")},
    ).json()

    second = client.post(
        f"/api/books/{book['book_id']}/sources/upload",
        files={"file": ("a-renamed.txt", content, "text/plain")},
    ).json()

    assert first["source_id"] == second["source_id"]

    all_sources = client.get(f"/api/books/{book['book_id']}/sources").json()
    assert len(all_sources) == 1


def test_reanalyze_with_new_purpose_updates_same_profile_row(client):
    book = _create_book(client)
    source = client.post(
        f"/api/books/{book['book_id']}/sources/upload",
        files={"file": ("data.csv", b"a,b\n1,2\n3,4\n", "text/csv")},
    ).json()

    first = client.post(
        f"/api/sources/{source['source_id']}/analyze", json={"purpose": "일반 분석"}
    ).json()
    assert "일반 분석" in first["summary"]

    second = client.post(
        f"/api/sources/{source['source_id']}/analyze", json={"purpose": "전문가용 기술서"}
    ).json()
    assert "전문가용 기술서" in second["summary"]
    assert second["analysis_purpose"] == "전문가용 기술서"

    profile_resp = client.get(f"/api/sources/{source['source_id']}/profile").json()
    assert profile_resp["profile_id"] == first["profile_id"] == second["profile_id"]


def test_unsupported_file_type_rejected(client):
    book = _create_book(client)
    resp = client.post(
        f"/api/books/{book['book_id']}/sources/upload",
        files={"file": ("archive.zip", b"binary", "application/zip")},
    )
    assert resp.status_code == 400


def test_upload_docx_and_analyze_extracts_paragraphs_and_tables(client):
    from io import BytesIO

    from docx import Document

    doc = Document()
    doc.add_heading("ALD 공정 개요", level=1)
    doc.add_paragraph("이 문서는 원자층 증착 공정을 설명한다.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "parameter"
    table.rows[0].cells[1].text = "importance"
    table.rows[1].cells[0].text = "temperature"
    table.rows[1].cells[1].text = "high"

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    book = client.post("/api/books", json={"workspace_id": "ws-1", "title": "테스트 책"}).json()
    resp = client.post(
        f"/api/books/{book['book_id']}/sources/upload",
        files={
            "file": (
                "note.docx",
                buf.read(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert resp.status_code == 200, resp.text
    source = resp.json()
    assert source["source_type"] == "docx"

    analyze_resp = client.post(f"/api/sources/{source['source_id']}/analyze", json={})
    assert analyze_resp.status_code == 200, analyze_resp.text

    source_after = client.get(f"/api/books/{book['book_id']}/sources").json()[0]
    assert source_after["status"] == "analyzed"
    assert "원자층 증착 공정" in source_after["raw_text"]
    assert "temperature" in source_after["raw_text"]
    assert "high" in source_after["raw_text"]


def test_extract_url_parses_title_and_strips_html():
    html = """
    <html><head><title>ALD 공정 가이드</title></head>
    <body><script>ignore me</script><p>온도와 압력이 중요합니다.</p></body></html>
    """
    title, text = extract_url("https://example.com/ald", http_get=lambda url: html)
    assert title == "ALD 공정 가이드"
    assert "ignore me" not in text
    assert "온도와 압력이 중요합니다" in text


def test_register_local_dir_registers_and_auto_analyzes(client, tmp_path):
    (tmp_path / "report1.txt").write_text("사출기 온도 데이터입니다.", encoding="utf-8")
    (tmp_path / "report2.md").write_text("# 사출기 압력 데이터", encoding="utf-8")
    (tmp_path / "image.png").write_bytes(b"not a real image")
    sub = tmp_path / "sub"
    sub.mkdir()
    (sub / "report3.txt").write_text("서브 폴더 자료", encoding="utf-8")

    book = _create_book(client)

    resp = client.post(
        f"/api/books/{book['book_id']}/sources/register-local-dir",
        json={"dir_path": str(tmp_path)},
    )
    assert resp.status_code == 200, resp.text
    body = resp.json()

    registered_titles = {r["title"] for r in body["registered"]}
    assert registered_titles == {"report1.txt", "report2.md"}
    assert body["skipped"] == ["image.png"]
    assert body["failed"] == []
    assert all(r["status"] == "analyzed" for r in body["registered"])

    sources_after = client.get(f"/api/books/{book['book_id']}/sources").json()
    assert len(sources_after) == 2
    assert all(s["status"] == "analyzed" for s in sources_after)


def test_register_local_dir_recursive_includes_subfolders(client, tmp_path):
    (tmp_path / "report1.txt").write_text("자료1", encoding="utf-8")
    sub = tmp_path / "sub"
    sub.mkdir()
    (sub / "report2.txt").write_text("자료2", encoding="utf-8")

    book = _create_book(client)

    resp = client.post(
        f"/api/books/{book['book_id']}/sources/register-local-dir",
        json={"dir_path": str(tmp_path), "recursive": True},
    )
    registered_titles = {r["title"] for r in resp.json()["registered"]}
    assert registered_titles == {"report1.txt", "report2.txt"}


def test_register_local_dir_rerun_does_not_reanalyze(client, tmp_path):
    (tmp_path / "report1.txt").write_text("자료1", encoding="utf-8")
    book = _create_book(client)

    client.post(
        f"/api/books/{book['book_id']}/sources/register-local-dir",
        json={"dir_path": str(tmp_path)},
    )

    call_count = {"n": 0}

    def _counting_llm(system_prompt: str, user_prompt: str) -> str:
        call_count["n"] += 1
        return _fake_llm_call(system_prompt, user_prompt)

    app.dependency_overrides[get_llm_call] = lambda: _counting_llm

    resp = client.post(
        f"/api/books/{book['book_id']}/sources/register-local-dir",
        json={"dir_path": str(tmp_path)},
    )
    assert resp.status_code == 200, resp.text
    assert call_count["n"] == 0
