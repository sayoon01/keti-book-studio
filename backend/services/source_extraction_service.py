from __future__ import annotations

import csv
import json
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from pypdf import PdfReader


class UnsupportedSourceTypeError(ValueError):
    """현재 추출기가 지원하지 않는 파일 형식입니다."""


class SourceExtractionError(RuntimeError):
    """파일 추출 과정에서 발생한 오류입니다."""


@dataclass
class ExtractedSourceContent:
    text: str
    source_type: str
    title: str

    character_count: int
    row_count: int | None = None
    column_count: int | None = None

    metadata: dict[str, Any] | None = None


class SourceExtractionService:
    SUPPORTED_EXTENSIONS = {
        ".txt",
        ".md",
        ".pdf",
        ".docx",
        ".csv",
        ".xlsx",
        ".xls",
        ".json",
        ".yaml",
        ".yml",
        ".py",
        ".js",
        ".jsx",
        ".ts",
        ".tsx",
        ".sql",
        ".html",
        ".css",
    }

    def extract(self, file_path: Path) -> ExtractedSourceContent:
        if not file_path.exists():
            raise SourceExtractionError(
                f"파일을 찾을 수 없습니다: {file_path}"
            )

        extension = file_path.suffix.lower()

        try:
            if extension in {
                ".txt",
                ".md",
                ".yaml",
                ".yml",
                ".py",
                ".js",
                ".jsx",
                ".ts",
                ".tsx",
                ".sql",
                ".html",
                ".css",
            }:
                return self._extract_text_file(file_path)

            if extension == ".pdf":
                return self._extract_pdf(file_path)

            if extension == ".docx":
                return self._extract_docx(file_path)

            if extension == ".csv":
                return self._extract_csv(file_path)

            if extension in {".xlsx", ".xls"}:
                return self._extract_excel(file_path)

            if extension == ".json":
                return self._extract_json(file_path)

        except UnsupportedSourceTypeError:
            raise

        except Exception as exc:
            raise SourceExtractionError(
                f"파일 추출에 실패했습니다: {file_path.name}: {exc}"
            ) from exc

        raise UnsupportedSourceTypeError(
            f"지원하지 않는 파일 형식입니다: {extension}"
        )

    def _extract_text_file(
        self,
        file_path: Path,
    ) -> ExtractedSourceContent:
        text = self._read_text_with_fallback(file_path)

        return ExtractedSourceContent(
            text=text,
            source_type=file_path.suffix.lower().lstrip("."),
            title=file_path.stem,
            character_count=len(text),
            metadata={
                "extension": file_path.suffix.lower(),
            },
        )

    def _extract_pdf(
        self,
        file_path: Path,
    ) -> ExtractedSourceContent:
        reader = PdfReader(str(file_path))

        pages: list[str] = []

        for page_index, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text() or ""

            pages.append(
                f"\n\n--- Page {page_index} ---\n\n{page_text}"
            )

        text = "".join(pages).strip()

        return ExtractedSourceContent(
            text=text,
            source_type="pdf",
            title=file_path.stem,
            character_count=len(text),
            metadata={
                "page_count": len(reader.pages),
            },
        )

    def _extract_docx(
        self,
        file_path: Path,
    ) -> ExtractedSourceContent:
        document = Document(str(file_path))

        paragraphs = [
            paragraph.text
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]

        table_texts: list[str] = []

        for table_index, table in enumerate(
            document.tables,
            start=1,
        ):
            rows: list[str] = []

            for row in table.rows:
                row_values = [
                    cell.text.strip()
                    for cell in row.cells
                ]
                rows.append(" | ".join(row_values))

            table_texts.append(
                f"\n\n[표 {table_index}]\n"
                + "\n".join(rows)
            )

        text = "\n\n".join(paragraphs) + "".join(table_texts)

        return ExtractedSourceContent(
            text=text.strip(),
            source_type="docx",
            title=file_path.stem,
            character_count=len(text),
            metadata={
                "paragraph_count": len(paragraphs),
                "table_count": len(document.tables),
            },
        )

    def _extract_csv(
        self,
        file_path: Path,
    ) -> ExtractedSourceContent:
        dataframe = self._read_csv_with_fallback(file_path)

        preview = dataframe.head(200)

        text_parts = [
            f"# CSV 파일: {file_path.name}",
            "",
            f"- 전체 행 수: {len(dataframe)}",
            f"- 전체 열 수: {len(dataframe.columns)}",
            f"- 컬럼: {', '.join(map(str, dataframe.columns))}",
            "",
            "## 미리보기",
            preview.to_csv(index=False),
        ]

        text = "\n".join(text_parts)

        return ExtractedSourceContent(
            text=text,
            source_type="csv",
            title=file_path.stem,
            character_count=len(text),
            row_count=len(dataframe),
            column_count=len(dataframe.columns),
            metadata={
                "columns": list(map(str, dataframe.columns)),
                "preview_rows": len(preview),
            },
        )

    def _extract_excel(
        self,
        file_path: Path,
    ) -> ExtractedSourceContent:
        workbook = pd.ExcelFile(file_path)

        text_parts = [
            f"# Excel 파일: {file_path.name}",
            f"- 시트 수: {len(workbook.sheet_names)}",
            f"- 시트: {', '.join(workbook.sheet_names)}",
        ]

        total_rows = 0
        maximum_columns = 0

        sheets_metadata: list[dict[str, Any]] = []

        for sheet_name in workbook.sheet_names:
            dataframe = pd.read_excel(
                file_path,
                sheet_name=sheet_name,
            )

            total_rows += len(dataframe)
            maximum_columns = max(
                maximum_columns,
                len(dataframe.columns),
            )

            preview = dataframe.head(100)

            text_parts.extend(
                [
                    "",
                    f"## 시트: {sheet_name}",
                    f"- 행 수: {len(dataframe)}",
                    f"- 열 수: {len(dataframe.columns)}",
                    f"- 컬럼: {', '.join(map(str, dataframe.columns))}",
                    "",
                    preview.to_csv(index=False),
                ]
            )

            sheets_metadata.append(
                {
                    "sheet_name": sheet_name,
                    "row_count": len(dataframe),
                    "column_count": len(dataframe.columns),
                    "columns": list(
                        map(str, dataframe.columns)
                    ),
                }
            )

        text = "\n".join(text_parts)

        return ExtractedSourceContent(
            text=text,
            source_type=file_path.suffix.lower().lstrip("."),
            title=file_path.stem,
            character_count=len(text),
            row_count=total_rows,
            column_count=maximum_columns,
            metadata={
                "sheets": sheets_metadata,
            },
        )

    def _extract_json(
        self,
        file_path: Path,
    ) -> ExtractedSourceContent:
        raw_text = self._read_text_with_fallback(file_path)
        parsed = json.loads(raw_text)

        formatted = json.dumps(
            parsed,
            ensure_ascii=False,
            indent=2,
        )

        return ExtractedSourceContent(
            text=formatted,
            source_type="json",
            title=file_path.stem,
            character_count=len(formatted),
            metadata={
                "root_type": type(parsed).__name__,
            },
        )

    def _read_text_with_fallback(
        self,
        file_path: Path,
    ) -> str:
        encodings = [
            "utf-8",
            "utf-8-sig",
            "cp949",
            "euc-kr",
        ]

        last_error: Exception | None = None

        for encoding in encodings:
            try:
                return file_path.read_text(
                    encoding=encoding,
                )
            except UnicodeDecodeError as exc:
                last_error = exc

        raise SourceExtractionError(
            f"텍스트 인코딩을 판별할 수 없습니다: {last_error}"
        )

    def _read_csv_with_fallback(
        self,
        file_path: Path,
    ) -> pd.DataFrame:
        encodings = [
            "utf-8",
            "utf-8-sig",
            "cp949",
            "euc-kr",
        ]

        last_error: Exception | None = None

        for encoding in encodings:
            try:
                return pd.read_csv(
                    file_path,
                    encoding=encoding,
                    low_memory=False,
                )
            except UnicodeDecodeError as exc:
                last_error = exc

        raise SourceExtractionError(
            f"CSV 인코딩을 판별할 수 없습니다: {last_error}"
        )
